import io
import logging
from collections import defaultdict

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _set_on_off(parent, tag: str, enabled: bool = True) -> OxmlElement:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:val"), "1" if enabled else "0")
    return element


def _set_paragraph_rtl(paragraph) -> None:
    """
    Word bidi paragraphs behave correctly when jc is set to left.
    This is visually right-aligned in RTL documents.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    _set_on_off(p_pr, "w:bidi", True)

    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)

    jc.set(qn("w:val"), "left")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_run_rtl(run, font_name: str = "B Nazanin") -> None:
    r_pr = run._element.get_or_add_rPr()
    _set_on_off(r_pr, "w:rtl", True)
    _set_on_off(r_pr, "w:cs", True)

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:bidi"), "fa-IR")

    if font_name:
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "1")


def _apply_table_rtl(table) -> None:
    tbl_pr = table._tbl.tblPr

    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi_visual)
    bidi_visual.set(qn("w:val"), "1")


def _safe_heading_level(value, default: int = 1) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        level = default
    return max(1, min(level, 9))


class DocumentExportService:
    def __init__(self, document):
        self.document = document

    def export_normalized_docx(self) -> bytes:
        doc = DocxDocument()
        all_blocks = list(self.document.blocks.all().order_by("order_index"))

        # Group table blocks by table_index, regardless of block_type
        tables_map = defaultdict(list)
        for b in all_blocks:
            # For raw_text, only blocks with complete table coordinates should be treated as table blocks
            if (
                b.table_index is not None
                and b.row_index is not None
                and b.cell_index is not None
            ):
                tables_map[b.table_index].append(b)

        # Sort blocks inside each table
        for t_idx in list(tables_map.keys()):
            tables_map[t_idx].sort(
                key=lambda x: (
                    x.row_index or 0,
                    x.cell_index or 0,
                    x.cell_paragraph_index or 0,
                    x.order_index or 0,
                )
            )

        rendered_tables = set()

        for block in all_blocks:
            b_type = (block.block_type or "").lower()
            text = block.normalized_text or block.raw_text or ""
            is_rtl = block.is_rtl if block.is_rtl is not None else True
            align = block.alignment or "left"  # kept as in your code, even if unused

            # Table blocks are detected by table_index, not by block_type
            is_table_block = (
                block.table_index is not None
                and block.row_index is not None
                and block.cell_index is not None
            )

            if is_table_block:
                t_idx = block.table_index

                if t_idx in rendered_tables:
                    continue

                t_blocks = tables_map.get(t_idx, [])
                if not t_blocks:
                    continue

                first_block_of_table = min(t_blocks, key=lambda x: x.order_index or 0)
                if block.id != first_block_of_table.id:
                    continue

                max_r = max((tb.row_index or 0) for tb in t_blocks)
                max_c = max((tb.cell_index or 0) for tb in t_blocks)

                table = doc.add_table(rows=max_r + 1, cols=max_c + 1)
                table.style = "Table Grid"
                _apply_table_rtl(table)

                for row in table.rows:
                    _set_row_cant_split(row)

                cell_content = defaultdict(list)
                for tb in t_blocks:
                    cell_content[(tb.row_index or 0, tb.cell_index or 0)].append(tb)

                for (r, c), cell_blocks in cell_content.items():
                    cell = table.cell(r, c)

                    # Remove extra empty paragraphs created by python-docx
                    while len(cell.paragraphs) > 1:
                        p_to_remove = cell.paragraphs[-1]._p
                        p_to_remove.getparent().remove(p_to_remove)

                    first_para = True
                    for tb in cell_blocks:
                        cell_text = tb.normalized_text or tb.raw_text or ""

                        if first_para:
                            p = cell.paragraphs[0]
                            p.text = ""
                            first_para = False
                        else:
                            p = cell.add_paragraph()

                        run = p.add_run(cell_text)

                        # Critical: keep exactly the same RTL logic as your working code
                        _set_paragraph_rtl(p)
                        _set_run_rtl(run, font_name="B Nazanin")

                        # Optional spacing controls
                        p.paragraph_format.keep_together = True
                        p.paragraph_format.widow_control = True

                rendered_tables.add(t_idx)
                continue

            # Non-table content
            if b_type == "heading":
                level = _safe_heading_level(block.heading_level)
                p = doc.add_heading(text, level=level)
            else:
                p = doc.add_paragraph(text)

            if is_rtl:
                _set_paragraph_rtl(p)
                for run in p.runs:
                    _set_run_rtl(run, font_name="B Nazanin")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
