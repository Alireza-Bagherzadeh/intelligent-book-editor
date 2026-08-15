# doc_process/services/docx_parser.py
from __future__ import annotations

import logging
import re
import statistics
from pathlib import Path
from typing import Any, Dict, Iterator, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from doc_process.services.text_normalizer import TextNormalizationService

logger = logging.getLogger(__name__)

FA_TO_EN_DIGITS = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")


class DocxParseService:
    """
    Parse DOCX files, extract structural metadata, identify headings/paragraphs,
    resolve hierarchical parent relationships, and output a flat list of block schemas.
    """

    def __init__(self, include_empty_blocks: bool = False) -> None:
        self.include_empty_blocks = include_empty_blocks
        # Shared text normalization service
        self.normalizer = TextNormalizationService()

    def extract_blocks(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Parses the DOCX file, applies normalization, identifies headings,
        resolves parent headings, extracts rich page formatting, and returns a flat list of blocks.
        """
        if not file_path.exists():
            logger.error(f"[DocxParseService] File not found at path: {file_path}")
            raise FileNotFoundError(f"File not found at: {file_path}")

        try:
            doc = Document(str(file_path))
        except Exception as e:
            logger.error(f"[DocxParseService] Failed to open document with python-docx: {e}")
            raise

        body_size = self._body_size_median(doc, default=13.0)
        elements = list(self._load_docx_elements(doc))
        
        logger.info(
            f"[DocxParseService] Path: {file_path} | Found {len(doc.paragraphs)} raw paragraphs "
            f"and {len(doc.tables)} tables. Parsed elements: {len(elements)}"
        )

        blocks: List[Dict[str, Any]] = []
        active_headings: Dict[int, int] = {}
        first_content = True

        for item in elements:
            element = item["element"]
            text = getattr(element, "text", "") or ""
            raw_text = text
            # Use the shared normalizer instead of local methods
            normalized_text = self.normalizer.normalize(text)

            # Skip empty paragraphs if configured
            if not self.include_empty_blocks and not normalized_text.strip():
                continue

            style_name = self._style_name(element) if isinstance(element, Paragraph) else ""
            block_type = "paragraph"
            heading_level = None
            was_reclassified = False

            if isinstance(element, Paragraph):
                raw_semantic_level = self._heading_level_from_semantics(element)

                # Validation: Does the text resemble body text characteristics (length, sentences)?
                if self._looks_like_body(normalized_text):
                    # Reclassify to paragraph if styled as heading but contains body content
                    heading_level = None
                    block_type = "paragraph"
                    if raw_semantic_level is not None:
                        was_reclassified = True
                else:
                    heading_level = raw_semantic_level
                    if heading_level is None:
                        heading_level = self._looks_like_heading(
                            paragraph=element,
                            body_size=body_size,
                            first_content=first_content,
                        )

                    if heading_level is not None:
                        block_type = "heading"

            current_order_index = len(blocks)
            parent_heading_order_index = None

            if block_type == "heading" and heading_level is not None:
                parent_level = heading_level - 1
                if parent_level > 0 and parent_level in active_headings:
                    parent_heading_order_index = active_headings[parent_level]

                active_headings[heading_level] = current_order_index

                # Clear deeper nested headings
                deeper_levels = [lvl for lvl in active_headings if lvl > heading_level]
                for lvl in deeper_levels:
                    del active_headings[lvl]
            else:
                if active_headings:
                    deepest_level = max(active_headings.keys())
                    parent_heading_order_index = active_headings[deepest_level]

            blocks.append({
                "block_type": block_type,
                "heading_level": heading_level,
                "order_index": current_order_index,
                "parent_heading_order_index": parent_heading_order_index,
                "raw_text": raw_text,
                "normalized_text": normalized_text,
                "style_name": style_name,
                "is_rtl": self._check_rtl(text),
                "alignment": self._get_alignment(element),
                "paragraph_index": item.get("paragraph_index"),
                "table_index": item.get("table_index"),
                "row_index": item.get("row_index"),
                "cell_index": item.get("cell_index"),
                "cell_paragraph_index": item.get("cell_paragraph_index"),
                "source_path": str(file_path),
                "format_metadata": self._format_metadata(
                    element=element,
                    body_size=body_size,
                    was_reclassified=was_reclassified,
                ),
            })

            if normalized_text.strip():
                first_content = False

        logger.info(f"[DocxParseService] Successfully generated {len(blocks)} blocks from elements.")
        return blocks

    def _load_docx_elements(self, doc: Document) -> Iterator[Dict[str, Any]]:
        for p_idx, paragraph in enumerate(doc.paragraphs):
            yield {
                "kind": "paragraph",
                "element": paragraph,
                "paragraph_index": p_idx,
                "table_index": None,
                "row_index": None,
                "cell_index": None,
                "cell_paragraph_index": None,
            }

        for t_idx, table in enumerate(doc.tables):
            yield from self._iter_table_elements(table, table_index=t_idx)

    def _iter_table_elements(self, table: Table, table_index: int) -> Iterator[Dict[str, Any]]:
        seen_cells: set[int] = set()

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)

                for p_idx, paragraph in enumerate(cell.paragraphs):
                    yield {
                        "kind": "table_cell_paragraph",
                        "element": paragraph,
                        "paragraph_index": None,
                        "table_index": table_index,
                        "row_index": r_idx,
                        "cell_index": c_idx,
                        "cell_paragraph_index": p_idx,
                    }

                for nested_table in cell.tables:
                    yield from self._iter_table_elements(nested_table, table_index=table_index)

    def _style_name(self, paragraph: Paragraph) -> str:
        return (paragraph.style.name or "") if paragraph.style else ""

    def _heading_level_from_semantics(self, paragraph: Paragraph) -> int | None:
        level = self._heading_level_from_style(paragraph)
        if level is not None:
            return level

        p = paragraph._p
        ppr = p.pPr
        if ppr is not None and ppr.outlineLvl is not None and ppr.outlineLvl.val is not None:
            try:
                return int(ppr.outlineLvl.val) + 1
            except (TypeError, ValueError):
                pass
        return None

    def _heading_level_from_style(self, paragraph: Paragraph) -> int | None:
        name = self._style_name(paragraph).strip()
        name_lower = name.lower()

        if "title" in name_lower or "عنوان سند" in name:
            return 1
        match = re.search(r"heading\s*([1-9])", name_lower)
        if match:
            return int(match.group(1))
        match = re.search(r"عنوان\s*([1-9])", name)
        if match:
            return int(match.group(1))
        match = re.search(r"تیتر\s*([1-9])", name)
        if match:
            return int(match.group(1))
        return None

    def _resolve_run_font_size(self, run: Run, paragraph: Paragraph) -> float | None:
        """
        Extract font size with support for style inheritance hierarchies in Word.
        """
        if run.font and run.font.size:
            return run.font.size.pt

        style = paragraph.style
        while style:
            if hasattr(style, "font") and style.font and style.font.size:
                return style.font.size.pt
            style = getattr(style, "base_style", None)

        return None

    def _paragraph_max_size(self, paragraph: Paragraph, default_size: float = 12.0) -> float:
        """
        Calculate the effective maximum font size of a paragraph looking at Runs and base styles.
        """
        sizes = [self._resolve_run_font_size(run, paragraph) for run in paragraph.runs]
        valid_sizes = [s for s in sizes if s is not None]

        if valid_sizes:
            return max(valid_sizes)

        # Fallback to paragraph style base tree
        style = paragraph.style
        while style:
            if hasattr(style, "font") and style.font and style.font.size:
                return style.font.size.pt
            style = getattr(style, "base_style", None)

        return default_size

    def _paragraph_bold_ratio(self, paragraph: Paragraph) -> float:
        total = 0
        bold = 0
        for run in paragraph.runs:
            length = len((run.text or "").strip())
            total += length
            if run.bold:
                bold += length
        return bold / total if total else 0.0

    def _body_size_median(self, doc: Document, default: float) -> float:
        sizes: List[float] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if not text or self._heading_level_from_semantics(paragraph) is not None:
                continue
            for run in paragraph.runs:
                size = self._resolve_run_font_size(run, paragraph)
                if size:
                    sizes.append(size)
        return statistics.median(sizes) if sizes else default

    def _numbering_depth(self, text: str) -> int | None:
        start = text.translate(FA_TO_EN_DIGITS)
        match = re.match(r"^\s*(\d+(?:\.\d+){0,4})[\)\-–—.:]?\s+", start)
        if match:
            return min(3, match.group(1).count(".") + 1)
        if re.match(r"^\s*فصل\s+\S+", text):
            return 1
        if re.match(r"^\s*بخش\s+\S+", text):
            return 2
        if re.match(r"^\s*ماده\s+\d+", start):
            return 2
        return None

    def _looks_like_body(self, text: str) -> bool:
        """
        Determine if the text exhibits typical paragraph body characteristics
        (e.g., long text block, multiple sentence delimiters).
        """
        clean_text = text.strip()
        if not clean_text:
            return False
        return len(clean_text) > 140 or (len(clean_text) > 90 and bool(re.search(r"[.!؟؛]\s", clean_text)))

    def _looks_like_heading(self, paragraph: Paragraph, body_size: float, first_content: bool) -> int | None:
        text = (paragraph.text or "").strip()
        if not text or self._looks_like_body(text):
            return None

        depth = self._numbering_depth(text)
        if depth:
            return depth

        max_size = self._paragraph_max_size(paragraph, default_size=body_size)
        bold_ratio = self._paragraph_bold_ratio(paragraph)
        terminal = bool(re.search(r"[.!؟؛،]$", text))

        common_heading_words = {
            "مقدمه", "چکیده", "نتیجه‌گیری", "نتیجه گیری", "جمع‌بندی", 
            "جمع بندی", "پیشنهادها", "پیشنهادات", "بحث", "روش تحقیق", 
            "یافته‌ها", "یافته ها", "منابع", "ضمائم", "پیوست"
        }
        if text in common_heading_words:
            return 1

        if first_content and len(text) <= 100 and ((max_size and max_size >= body_size + 3) or bold_ratio >= 0.7):
            return 1
        if not terminal and max_size and max_size >= body_size + 3:
            return 1
        if not terminal and (bold_ratio >= 0.65 or (max_size and max_size >= body_size + 1.5)):
            return 2
        if len(text) <= 60 and not terminal and self._is_centered(paragraph) and bold_ratio >= 0.4:
            return 1
        return None

    def _check_rtl(self, text: str) -> bool:
        if not text:
            return False
        rtl_chars = re.findall(r"[\u0590-\u08FF]", text)
        ltr_chars = re.findall(r"[A-Za-z]", text)
        return len(rtl_chars) >= len(ltr_chars)

    def _get_alignment(self, element: Any) -> str:
        alignment = getattr(element, "alignment", None)
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "justify"
        if alignment == WD_ALIGN_PARAGRAPH.LEFT:
            return "left"
        return "unknown"

    def _is_centered(self, paragraph: Paragraph) -> bool:
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def _format_metadata(self, element: Any, body_size: float = 12.0, was_reclassified: bool = False) -> Dict[str, Any]:
        if not isinstance(element, Paragraph):
            return {}

        # Extract pagination layout metadata properties
        p_format = getattr(element, "paragraph_format", None)
        keep_together = getattr(p_format, "keep_together", False) if p_format else False
        keep_with_next = getattr(p_format, "keep_with_next", False) if p_format else False
        widow_control = getattr(p_format, "widow_control", True) if p_format else True

        resolved_font_size = self._paragraph_max_size(element, default_size=body_size)

        return {
            "bold_ratio": round(self._paragraph_bold_ratio(element), 2),
            "max_font_size": resolved_font_size,
            "text_length": len((element.text or "").strip()),
            "raw_style_name": self._style_name(element),
            "semantic_heading_level": self._heading_level_from_semantics(element),
            "was_reclassified": was_reclassified,
            "pagination": {
                "keep_together": keep_together,
                "keep_with_next": keep_with_next,
                "widow_control": widow_control,
            }
        }
